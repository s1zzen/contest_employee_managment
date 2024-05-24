import sqlite3
from openpyxl import Workbook
from docx import Document
from src.config import Configuration
import os


def export_data(self):
    cfg = Configuration()
    os.makedirs(os.path.dirname(cfg.output_dir), exist_ok=True)
    conn = sqlite3.connect(cfg.db_path)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT Employees.name, Employees.position, Branches.name AS branch_name
        FROM Employees
        LEFT JOIN Branches ON Employees.branch_id = Branches.branch_id
    ''')
    data = cursor.fetchall()
    conn.close()

    # Export to XLSX
    wb = Workbook()
    ws = wb.active
    ws.append(['Имя', 'Должность', 'Филиал'])
    for row in data:
        ws.append(row)
    wb.save(cfg.output_dir + 'employees.xlsx')

    # Export to DOCX
    doc = Document()
    doc.add_heading('Список сотрудников', 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя'
    hdr_cells[1].text = 'Должность'
    hdr_cells[2].text = 'Филиал'
    for row in data:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
    doc.save(cfg.output_dir + 'employees.docx')
